"""Tests for document import cleanup and workflow splitting."""

from io import BytesIO
from types import SimpleNamespace

import web.app as web_app
from docx import Document

from src.importers.document_parser import DocumentParser
from src.importers.content_extractor import ContentExtractor
from src.importers.workflow_detector import WorkflowDetector
from src.parser.nlp_parser import NLPParser
from web.app import SAMPLE_WORKFLOWS


def test_preprocess_for_parser_drops_reference_text_and_keeps_actions():
    text = """
    # Table of Contents
    Phase 1.......................... 2
    Phase 2.......................... 3

    # Purpose
    This manual describes the full process in detail.

    ## Phase 1: Intake
    Procedure:
    1) Receive the device from the customer
    2) Verify the serial number
    Next Step:
    If approved, move to Repair

    ## Phase 2: Repair
    - Replace the failed component
    - Test the repaired unit
    """

    extractor = ContentExtractor()

    processed = extractor.preprocess_for_parser(text)

    assert "Table of Contents" not in processed
    assert "Purpose" not in processed
    assert "Procedure:" not in processed
    assert "Next Step:" not in processed
    assert "# Phase 1: Intake" in processed
    assert "# Phase 2: Repair" in processed
    assert "1. Receive the device from the customer" in processed
    assert "2. Verify the serial number" in processed
    assert "If approved, move to Repair" in processed
    assert "- Replace the failed component" in processed


def test_docx_heading_detection_treats_missing_style_as_body_text():
    parser = DocumentParser()
    paragraph = SimpleNamespace(
        style=None,
        text="Open the customer request and verify the account",
        runs=[SimpleNamespace(text="Open the customer request and verify the account", bold=False)],
    )

    assert parser._get_heading_level(paragraph) == 0


def test_docx_parser_formats_workflow_title_and_numbered_step_tables(tmp_path):
    doc_path = tmp_path / "workflow-table.docx"
    doc = Document()
    doc.add_paragraph("Client Training Manual")

    title_table = doc.add_table(rows=1, cols=2)
    title_table.cell(0, 0).text = "WF 1-A"
    title_table.cell(0, 1).text = "Intake Review"

    steps_table = doc.add_table(rows=3, cols=2)
    steps_table.cell(0, 0).text = "1"
    steps_table.cell(0, 1).text = "Open the request"
    steps_table.cell(1, 0).text = "2"
    steps_table.cell(1, 1).text = "Verify all required information"
    steps_table.cell(2, 0).text = "3"
    steps_table.cell(2, 1).text = "If complete, move the ticket to Ready"
    doc.save(doc_path)

    result = DocumentParser().parse(doc_path)

    assert result["success"] is True
    assert "## WF 1-A: Intake Review" in result["text"]
    assert "1. Open the request" in result["text"]
    assert "3. If complete, move the ticket to Ready" in result["text"]


def test_upload_docx_with_workflow_tables_detects_workflows():
    doc = Document()
    doc.add_paragraph("Training Workflow Pack")

    title_table = doc.add_table(rows=1, cols=2)
    title_table.cell(0, 0).text = "WF 2-B"
    title_table.cell(0, 1).text = "Order Fulfillment"

    steps_table = doc.add_table(rows=5, cols=2)
    rows = [
        ("1", "Receive the signed invoice"),
        ("2", "Create the customer order"),
        ("3", "Check whether payment is received"),
        ("4", "If payment is received, release the shipment"),
        ("5", "Send the tracking number and close the order"),
    ]
    for row_index, (number, text) in enumerate(rows):
        steps_table.cell(row_index, 0).text = number
        steps_table.cell(row_index, 1).text = text

    payload = BytesIO()
    doc.save(payload)
    payload.seek(0)

    with web_app.app.test_client() as client:
        response = client.post(
            "/api/upload",
            data={"file": (payload, "order-workflow.docx")},
            content_type="multipart/form-data",
        )

    assert response.status_code == 200
    body = response.get_json()
    assert body["success"] is True
    assert body["metadata"]["filename"] == "order-workflow.docx"
    assert body["metadata"]["tables"] == 2
    assert body["workflows"]
    assert body["summary"]["total_workflows"] >= 1


def test_content_extractor_keeps_user_login_sample_as_single_workflow():
    sample = SAMPLE_WORKFLOWS["user-login"]["text"]
    extractor = ContentExtractor()

    workflows = extractor.extract_workflows(sample)
    processed = extractor.preprocess_for_parser(workflows[0]["content"])
    summary = extractor.get_workflow_summary(processed)

    assert len(workflows) == 1
    assert "1. User opens login page" in processed
    assert "5. If 2FA enabled, send verification code" in processed
    assert "9. End" in processed
    assert summary["step_count"] == 9


def test_enabled_does_not_make_action_line_a_header():
    extractor = ContentExtractor()

    assert extractor._is_header("If 2FA enabled, send verification code") is False
    assert extractor._is_header("5. If 2FA enabled, send verification code") is False


def test_auto_detect_keeps_manual_sections_separate():
    text = """
    # 1. Product Knowledge
    1. Review product lineup
    2. Confirm supported models
    3. Document key differences

    # 2. Intake Review
    1. Open the service request
    2. Validate the serial number
    3. Record the failure notes

    # 3. Fulfillment
    1. Pack the repaired unit
    2. Send the shipment notification
    3. Close the request
    """

    detector = WorkflowDetector(split_mode="auto")

    workflows = detector.detect_workflows(text)

    assert len(workflows) == 3
    assert [workflow.title for workflow in workflows] == [
        "1. Product Knowledge",
        "2. Intake Review",
        "3. Fulfillment",
    ]


def test_auto_detect_merges_phased_stage_sections():
    text = """
    # Stage 1: New Request
    1. Receive the request
    2. Validate required fields
    3. Move the ticket to "Queued"

    # Stage 2: Triage
    1. Review the request details
    2. Assign the owner
    3. Move the ticket to "Repair"

    # Stage 3: Repair
    1. Repair the device
    2. Test the output
    3. Move the ticket to "Ready to Ship"

    # Stage 4: Closeout
    1. Confirm payment
    2. Ship the device
    3. Update ticket to "Closed"
    """

    detector = WorkflowDetector(split_mode="auto")

    workflows = detector.detect_workflows(text)

    assert len(workflows) == 1
    assert workflows[0].title == "Pipeline: Stage 1: New Request to Stage 4: Closeout"
    assert "## Stage 1: New Request" in workflows[0].content
    assert "## Stage 4: Closeout" in workflows[0].content


def test_auto_detect_keeps_short_stage_sections():
    text = """
    # Stage 1: Intake
    Once the order is received, create the intake record

    # Stage 2: Ready
    Move the ticket to "Ready"
    """

    detector = WorkflowDetector(split_mode="auto")

    workflows = detector.detect_workflows(text)

    assert len(workflows) == 2
    assert [workflow.title for workflow in workflows] == [
        "Stage 1: Intake",
        "Stage 2: Ready",
    ]


def test_section_filter_skips_prerequisites_heading():
    text = """
    # Prerequisites
    1. Confirm admin access
    2. Download the support image

    # Device Setup
    1. Boot the device
    2. Install the image
    3. Verify connectivity
    """

    detector = WorkflowDetector(split_mode="section")

    workflows = detector.detect_workflows(text)

    assert len(workflows) == 1
    assert workflows[0].title == "Device Setup"


def test_merged_pipeline_text_keeps_phase_groups_for_export():
    text = """
    # 1. Intake
    1. Open the repair request
    If all required information is present:
    - Move the ticket to 'Label Sent to Customer'

    # 2. Label Sent to Customer
    - Monitor the tracking ID
    If tracking shows the package has shipped:
    - Move the ticket to 'Shipped'

    # 3. Shipped
    - Confirm delivery
    """

    detector = WorkflowDetector(split_mode="auto")
    sections = detector._try_header_detection(text.strip().splitlines())
    workflow = detector._merge_sections(detector._analyze_and_filter(sections))

    steps = NLPParser(use_spacy=False).parse(workflow.content)

    assert steps
    assert all(not step.text.startswith("##") for step in steps)
    assert [step.group for step in steps[:4]] == [
        "1. Intake",
        "1. Intake",
        "2. Label Sent to Customer",
        "2. Label Sent to Customer",
    ]
