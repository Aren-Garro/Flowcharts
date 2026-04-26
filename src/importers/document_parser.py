"""Multi-format document parser for extracting text from various file types."""

import logging
import re
from pathlib import Path
from typing import Any, Dict, List

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse various document formats to extract raw text."""

    def __init__(self):
        self.supported_formats = [".txt", ".md", ".pdf", ".docx", ".doc"]
        self._check_dependencies()

    def _check_dependencies(self):
        self.has_pdf = False
        self.has_docx = False

        try:
            __import__("PyPDF2")
            self.has_pdf = True
            logger.info("PDF support enabled (PyPDF2)")
        except ImportError:
            try:
                __import__("pdfplumber")
                self.has_pdf = True
                logger.info("PDF support enabled (pdfplumber)")
            except ImportError:
                logger.warning("PDF support not available. Install PyPDF2 or pdfplumber.")

        try:
            __import__("docx")
            self.has_docx = True
            logger.info("DOCX support enabled (python-docx)")
        except ImportError:
            logger.warning("DOCX support not available. Install python-docx.")

    def parse(self, file_path: Path, encoding: str = "utf-8") -> Dict[str, Any]:
        file_path = Path(file_path)

        if not file_path.exists():
            return {"text": "", "metadata": {}, "format": None, "success": False,
                    "error": f"File not found: {file_path}"}

        suffix = file_path.suffix.lower()
        if suffix not in self.supported_formats:
            return {"text": "", "metadata": {}, "format": suffix, "success": False,
                    "error": f"Unsupported format: {suffix}"}

        try:
            if suffix in [".txt", ".md"]:
                return self._parse_text(file_path, encoding)
            elif suffix == ".pdf":
                return self._parse_pdf(file_path)
            elif suffix in [".docx", ".doc"]:
                return self._parse_docx(file_path)
            else:
                return {"text": "", "metadata": {}, "format": suffix, "success": False,
                        "error": f"Format not yet implemented: {suffix}"}
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            return {"text": "", "metadata": {}, "format": suffix, "success": False,
                    "error": str(e)}

    def _parse_text(self, file_path: Path, encoding: str) -> Dict[str, Any]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read()
            return {"text": text, "metadata": {"filename": file_path.name,
                    "size": file_path.stat().st_size}, "format": file_path.suffix, "success": True}
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read()
            return {"text": text, "metadata": {"filename": file_path.name,
                    "size": file_path.stat().st_size, "encoding": "latin-1"},
                    "format": file_path.suffix, "success": True}

    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        if not self.has_pdf:
            return {"text": "", "metadata": {}, "format": ".pdf", "success": False,
                    "error": "PDF support not available. Install: pip install PyPDF2 pdfplumber"}
        try:
            __import__("pdfplumber")
            return self._parse_pdf_pdfplumber(file_path)
        except ImportError:
            pass
        try:
            __import__("PyPDF2")
            return self._parse_pdf_pypdf2(file_path)
        except ImportError:
            return {"text": "", "metadata": {}, "format": ".pdf", "success": False,
                    "error": "No PDF library available"}

    def _parse_pdf_pdfplumber(self, file_path: Path) -> Dict[str, Any]:
        import pdfplumber
        text_parts = []
        metadata = {}
        with pdfplumber.open(file_path) as pdf:
            metadata["pages"] = len(pdf.pages)
            metadata["filename"] = file_path.name
            if pdf.metadata:
                metadata["title"] = pdf.metadata.get("Title", "")
                metadata["author"] = pdf.metadata.get("Author", "")
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = "\n\n".join(text_parts)
        if not text.strip():
            return {"text": "", "metadata": metadata, "format": ".pdf", "success": False,
                    "error": "No extractable text found in PDF. Scanned/image-only PDFs require OCR before import."}
        return {"text": text, "metadata": metadata, "format": ".pdf", "success": True}

    def _parse_pdf_pypdf2(self, file_path: Path) -> Dict[str, Any]:
        import PyPDF2
        text_parts = []
        metadata = {}
        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            metadata["pages"] = len(pdf_reader.pages)
            metadata["filename"] = file_path.name
            if pdf_reader.metadata:
                metadata["title"] = pdf_reader.metadata.get("/Title", "")
                metadata["author"] = pdf_reader.metadata.get("/Author", "")
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = "\n\n".join(text_parts)
        if not text.strip():
            return {"text": "", "metadata": metadata, "format": ".pdf", "success": False,
                    "error": "No extractable text found in PDF. Scanned/image-only PDFs require OCR before import."}
        return {"text": text, "metadata": metadata, "format": ".pdf", "success": True}

    def _get_heading_level(self, paragraph) -> int:
        """Detect heading level from paragraph style.

        Returns 0 for body text, 1-3 for heading levels.
        """
        style = getattr(paragraph, "style", None)
        style_name = (getattr(style, "name", "") or "").lower()

        # Explicit heading styles
        if 'heading 1' in style_name or style_name == 'title':
            return 1
        if 'heading 2' in style_name or style_name == 'subtitle':
            return 2
        if 'heading 3' in style_name or 'heading 4' in style_name:
            return 3

        # Check for TOC styles (not headings)
        if 'toc' in style_name:
            return 0

        # Detect bold-only short paragraphs as implicit headings
        text = paragraph.text.strip()
        if text and len(text) < 80 and len(text) > 3:
            all_bold = all(
                run.bold for run in paragraph.runs
                if run.text.strip()
            ) if paragraph.runs else False
            if all_bold and not re.match(r'^\d+\.?\s', text):
                # Bold, short, and not a numbered step → likely a heading
                return 2

        return 0

    # Symbols that decorate decision/outcome rows in many SOP templates.
    _DECORATIVE_SYMBOLS = '✓✔✗✘✖◆◇▶▷❖➔→❓❔⚑⯕•●○▪▫◉'
    _SYMBOL_RE = re.compile(f'[{_DECORATIVE_SYMBOLS}]+')
    _LEADING_DECORATION_RE = re.compile(rf'^\s*([{_DECORATIVE_SYMBOLS}\-–—:>\s]+)')
    _DECISION_CUE_RE = re.compile(
        r'^\s*(?:decision|decide|check|question|verify|evaluate|approval|approve)\b[:\-\s]*',
        re.IGNORECASE,
    )
    _BINARY_OUTCOME_VOCAB = {
        'yes': 'Yes', 'no': 'No',
        'approved': 'Yes', 'rejected': 'No',
        'approve': 'Yes', 'reject': 'No',
        'accept': 'Yes', 'decline': 'No',
        'pass': 'Yes', 'fail': 'No',
        'success': 'Yes', 'failure': 'No', 'failed': 'No',
        'valid': 'Yes', 'invalid': 'No',
        'true': 'Yes', 'false': 'No',
        'ok': 'Yes', 'okay': 'Yes', 'not ok': 'No',
        'allow': 'Yes', 'deny': 'No',
        'continue': 'Yes', 'stop': 'No',
    }

    def _strip_decoration(self, text: str) -> str:
        """Remove leading symbols, dashes, and the matching outcome-vocab word."""
        text = self._LEADING_DECORATION_RE.sub('', text or '').strip()
        return self._SYMBOL_RE.sub('', text).strip(' -:–—\t')

    def _classify_outcome_token(self, side: str):
        """Return the canonical outcome label (Yes/No/...) and the cleaned remainder."""
        cleaned = self._strip_decoration(side)
        if not cleaned:
            return None, ''
        lowered = cleaned.lower()
        for key, label in self._BINARY_OUTCOME_VOCAB.items():
            if lowered == key or lowered.startswith(key + ' ') or lowered.startswith(key + ':') or lowered.startswith(key + '—') or lowered.startswith(key + '-'):
                remainder = cleaned[len(key):].lstrip(' :-–—\t').strip()
                return label, remainder
        return None, cleaned

    def _try_decision_marker_row(self, only: str):
        """Detect 1-cell decision marker / question rows and return a normalized line, or None."""
        cleaned = self._strip_decoration(only)
        if not cleaned:
            return None
        cue_match = self._DECISION_CUE_RE.match(cleaned)
        if cue_match:
            remainder = cleaned[cue_match.end():].strip()
            if remainder:
                if not remainder.endswith('?'):
                    remainder = remainder.rstrip('.') + '?'
                return f"\n? {remainder}\n"
            return None
        if cleaned.endswith('?') and len(cleaned) <= 200 and len(cleaned.split()) >= 2:
            return f"\n? {cleaned}\n"
        return None

    def _try_binary_outcome_row(self, only: str):
        """Detect 1-cell binary-outcome rows (Yes | No, Approved / Rejected, ...) and return paired lines."""
        if not only:
            return None
        # Choose the separator that splits the row into exactly two non-trivial sides.
        for sep in ('|', '\t'):
            if sep in only:
                parts = [p.strip() for p in only.split(sep)]
                if len(parts) == 2 and all(parts):
                    return self._format_outcome_pair(parts[0], parts[1])
        # Slash separator only when both sides classify as outcome vocab to avoid
        # mangling URLs and ordinary "X/Y" prose.
        if '/' in only and len(only) <= 200:
            parts = [p.strip() for p in only.split('/')]
            if len(parts) == 2 and all(parts):
                left_label, _ = self._classify_outcome_token(parts[0])
                right_label, _ = self._classify_outcome_token(parts[1])
                if left_label and right_label and left_label != right_label:
                    return self._format_outcome_pair(parts[0], parts[1])
        return None

    def _format_outcome_pair(self, left: str, right: str):
        """Emit two `Yes: ...` / `No: ...` lines (or matching pair) for a binary-outcome row."""
        left_label, left_rest = self._classify_outcome_token(left)
        right_label, right_rest = self._classify_outcome_token(right)
        if left_label and right_label and left_label != right_label:
            ordered = sorted(
                [(left_label, left_rest), (right_label, right_rest)],
                key=lambda pair: 0 if pair[0] == 'Yes' else 1,
            )
            return '\n'.join(f"{label}: {rest}".rstrip(': ') for label, rest in ordered if (label or rest))
        # Fall back: only emit if at least one side is a known outcome word.
        if left_label or right_label:
            yes_side = left_rest if left_label == 'Yes' else (right_rest if right_label == 'Yes' else None)
            no_side = left_rest if left_label == 'No' else (right_rest if right_label == 'No' else None)
            lines = []
            if yes_side is not None:
                lines.append(f"Yes: {yes_side}".rstrip(': '))
            if no_side is not None:
                lines.append(f"No: {no_side}".rstrip(': '))
            return '\n'.join(lines) if lines else None
        return None

    def _format_table_row(self, cells_text, row_idx: int) -> str:
        """Format a table row, detecting Step/Action, decision, and binary-outcome rows."""
        cells = [re.sub(r'\s+', ' ', c.strip()) for c in cells_text]

        # Skip empty rows
        if not any(cells):
            return ''

        first = cells[0].strip()
        first_lower = first.lower()

        # Skip header rows (Step | Action, etc.)
        if first_lower in ('step', '#', 'no', 'no.', 'number'):
            return ''

        # Workflow title tables often use a short ID cell followed by the workflow name.
        if re.match(r'^(wf|workflow)\s*[\w.-]+$', first_lower) and len(cells) >= 2:
            title = ' '.join(cells[1:]).strip()
            if title:
                return f"\n## {first}: {title}\n"

        if len(cells) == 1:
            only = cells[0]
            if re.match(r'^(section|phase|stage)\s+\d+\b', only, re.IGNORECASE):
                return f"\n# {only}\n"
            decision_line = self._try_decision_marker_row(only)
            if decision_line is not None:
                return decision_line
            outcome_lines = self._try_binary_outcome_row(only)
            if outcome_lines is not None:
                return outcome_lines
            return only

        # If first cell is a number, format as a numbered step
        if first.isdigit() and len(cells) >= 2:
            step_num = first
            action = ' '.join(cells[1:]).strip()
            if action:
                return f"{step_num}. {action}"

        # Two-cell rows where both sides are binary-outcome tokens (e.g., "Yes | No outcome split into columns")
        if len(cells) == 2:
            outcome_lines = self._format_outcome_pair(cells[0], cells[1])
            if outcome_lines is not None:
                return outcome_lines

        # Otherwise join with pipes
        return ' | '.join(cells)

    def _append_paragraph_text(self, text_parts: List[str], paragraph) -> None:
        text = paragraph.text.strip()
        if not text:
            return
        heading_level = self._get_heading_level(paragraph)
        if heading_level > 0:
            prefix = '#' * heading_level
            text_parts.append(f"\n{prefix} {text}\n")
            return
        text_parts.append(text)

    def _append_table_text(self, text_parts: List[str], table) -> None:
        text_parts.append('')
        for row_i, row in enumerate(table.rows):
            cells = [cell.text for cell in row.cells]
            formatted = self._format_table_row(cells, row_i)
            if formatted:
                text_parts.append(formatted)
        text_parts.append('')

    def _parse_docx_in_order(self, doc, text_parts: List[str]) -> None:
        para_idx = 0
        table_idx = 0

        from docx.oxml.ns import qn

        for elem in list(doc.element.body):
            if elem.tag == qn('w:p'):
                if para_idx < len(doc.paragraphs):
                    self._append_paragraph_text(text_parts, doc.paragraphs[para_idx])
                para_idx += 1
                continue

            if elem.tag == qn('w:tbl'):
                if table_idx < len(doc.tables):
                    self._append_table_text(text_parts, doc.tables[table_idx])
                table_idx += 1

    def _parse_docx_fallback(self, doc, text_parts: List[str]) -> None:
        for para in doc.paragraphs:
            self._append_paragraph_text(text_parts, para)
        for table in doc.tables:
            self._append_table_text(text_parts, table)

    def _extract_docx_metadata(self, doc, file_path: Path) -> Dict[str, Any]:
        metadata = {
            "filename": file_path.name,
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }
        if hasattr(doc, "core_properties"):
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
        return metadata

    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX preserving heading hierarchy and table structure."""
        if not self.has_docx:
            return {"text": "", "metadata": {}, "format": ".docx", "success": False,
                    "error": "DOCX support not available. Install: pip install python-docx"}

        import docx

        try:
            doc = docx.Document(file_path)
            text_parts = []

            try:
                self._parse_docx_in_order(doc, text_parts)
            except Exception:
                self._parse_docx_fallback(doc, text_parts)

            text = "\n".join(text_parts)

            # Clean up excessive blank lines
            text = re.sub(r'\n{4,}', '\n\n\n', text)
            metadata = self._extract_docx_metadata(doc, file_path)

            return {"text": text, "metadata": metadata, "format": ".docx", "success": True}

        except Exception as e:
            return {"text": "", "metadata": {}, "format": ".docx", "success": False,
                    "error": f"Error parsing DOCX: {str(e)}"}

    def parse_clipboard(self) -> Dict[str, Any]:
        try:
            import pyperclip
            text = pyperclip.paste()
            return {"text": text, "metadata": {"source": "clipboard"},
                    "format": "clipboard", "success": True}
        except ImportError:
            return {"text": "", "metadata": {}, "format": "clipboard", "success": False,
                    "error": "Clipboard support not available. Install: pip install pyperclip"}
        except Exception as e:
            return {"text": "", "metadata": {}, "format": "clipboard", "success": False,
                    "error": str(e)}

    def get_supported_formats(self) -> list:
        formats = [".txt", ".md"]
        if self.has_pdf:
            formats.append(".pdf")
        if self.has_docx:
            formats.extend([".docx", ".doc"])
        return formats
